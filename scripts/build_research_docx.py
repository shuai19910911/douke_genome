#!/usr/bin/env python3
"""Build the formal LegumeGenomeFM research-design DOCX from Markdown.

Run with:
    uv run --no-project --with python-docx scripts/build_research_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "LEGUMEGENOMEFM_RESEARCH_DESIGN.md"
OUTPUT = ROOT / "docs" / "LegumeGenomeFM_Research_Design_and_Benchmark_20260722.docx"

NAVY = "17365D"
TEAL = "176B6B"
GOLD = "C6922C"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF6DF"
LIGHT_GREY = "F2F4F7"
MID_GREY = "D9E1E8"
DARK = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia="Noto Sans CJK SC", latin="Aptos", size=None, bold=None, color=None) -> None:
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    r_fonts.set(qn("w:ascii"), "Aptos")
    r_pr.extend((r_fonts, color, underline))
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|https?://[^\s)>]+)"
)


def add_inline(paragraph, text: str, default_size: float | None = None) -> None:
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size)
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, east_asia="Noto Sans CJK SC", latin="Consolas", size=(default_size or 10) - 0.5)
            run.font.color.rgb = RGBColor.from_string("8B1E3F")
        elif token.startswith("["):
            label, url = token[1:].split("](", 1)
            add_hyperlink(paragraph, label, url[:-1])
        else:
            add_hyperlink(paragraph, token.rstrip(".,;，。；"), token.rstrip(".,;，。；"))
            suffix = token[len(token.rstrip(".,;，。；")) :]
            if suffix:
                suffix_run = paragraph.add_run(suffix)
                set_run_font(suffix_run, size=default_size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=default_size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))
    set_run_font(run, size=9, color="667085")


def add_toc(document: Document) -> None:
    p = document.add_paragraph()
    p.style = document.styles["Title"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("目录")
    set_run_font(r, size=22, bold=True, color=NAVY)
    p2 = document.add_paragraph()
    run = p2.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "在Microsoft Word中按Ctrl+A后按F9更新目录与页码。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, placeholder, end))
    set_run_font(run, size=10, color="667085")
    document.add_page_break()


def setup_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18

    specs = {
        "Title": (24, NAVY, True, 12, 6),
        "Heading 1": (17, NAVY, True, 16, 8),
        "Heading 2": (14, TEAL, True, 12, 6),
        "Heading 3": (11.5, NAVY, True, 9, 4),
    }
    for name, (size, color, bold, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.page_break_before = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans CJK SC")
        style.paragraph_format.space_after = Pt(3)


def setup_section(section) -> None:
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.8)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("LegumeGenomeFM  |  研究设计·数据合同·基准评测")
    set_run_font(hr, size=8.5, color="667085")
    p_pr = hp._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), MID_GREY)
    border.append(bottom)
    p_pr.append(border)

    fp = section.footer.paragraphs[0]
    add_page_number(fp)


def add_cover(document: Document) -> None:
    document.add_paragraph("\n\n")
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = banner.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 300, 280, 300, 280)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LegumeGenomeFM")
    set_run_font(r, size=30, bold=True, color="FFFFFF")
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("豆科超长上下文基因组基础模型")
    set_run_font(r2, size=16, bold=True, color="DCEAF7")

    document.add_paragraph("\n")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("研究设计、数据合同与基准评测方案")
    set_run_font(r, size=22, bold=True, color=NAVY)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Research Design, Data Contract and Benchmark Protocol")
    set_run_font(r, size=12, color=TEAL)

    document.add_paragraph("\n")
    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("文档版本", "1.0"),
        ("证据截止", "2026-07-22"),
        ("候选模型", "LegumeGenomeFM-HierMamba"),
        ("候选参数", "314,669,504（设计公式；H20运行时未冻结）"),
        ("数据状态", "74-source QC完成；schema-2 release未生成"),
        ("性能状态", "正式预训练和全部下游任务均未运行"),
    ]
    for i, (key, value) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        set_cell_shading(c1, PALE_BLUE)
        for c in (c1, c2):
            set_cell_margins(c, 90, 120, 90, 120)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = p1.add_run(key)
        set_run_font(rr, size=9.5, bold=True, color=NAVY)
        p2 = c2.paragraphs[0]
        rr = p2.add_run(value)
        set_run_font(rr, size=9.5)

    document.add_paragraph("\n")
    warning = document.add_table(rows=1, cols=1)
    wc = warning.cell(0, 0)
    set_cell_shading(wc, PALE_GOLD)
    set_cell_margins(wc, 160, 180, 160, 180)
    wp = wc.paragraphs[0]
    wr = wp.add_run("证据声明：截至本版日期，本模型没有正式预训练或下游性能。所有外部数值均为文献原协议结果；设计目标、数据QC与代码测试不得表述为模型成绩。")
    set_run_font(wr, size=10, bold=True, color="6E4C00")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    r = p.add_run("LegumeGenomeFM Project  ·  2026")
    set_run_font(r, size=10, color="667085")
    document.add_page_break()


def parse_table(lines: list[str], start: int):
    raw = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw.append(lines[i].strip())
        i += 1
    if len(raw) < 2:
        return None, start
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in raw]
    if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        return None, start
    return [rows[0], *rows[2:]], i


def add_table(document: Document, rows: list[list[str]]) -> None:
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    font_size = 7.3 if columns >= 6 else (8.1 if columns >= 5 else 8.7)
    for ri, row in enumerate(rows):
        tr_pr = table.rows[ri]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if ri == 0:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)
        for ci in range(columns):
            cell = table.cell(ri, ci)
            value = row[ci] if ci < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 60, 70, 60, 70)
            set_cell_shading(cell, NAVY if ri == 0 else (LIGHT_GREY if ri % 2 == 0 else "FFFFFF"))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, default_size=font_size)
            for run in p.runs:
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def render_markdown(document: Document, source: str) -> None:
    lines = source.splitlines()
    i = 0
    first_h1_skipped = False
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        table_rows, next_i = parse_table(lines, i)
        if table_rows is not None:
            add_table(document, table_rows)
            i = next_i
            continue
        if stripped.startswith("# "):
            if not first_h1_skipped:
                first_h1_skipped = True
            else:
                p = document.add_paragraph(style="Heading 1")
                add_inline(p, stripped[2:])
            i += 1
            continue
        if stripped.startswith("## "):
            p = document.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
            i += 1
            continue
        if stripped.startswith("### "):
            p = document.add_paragraph(style="Heading 2")
            add_inline(p, stripped[4:])
            i += 1
            continue
        if stripped.startswith("#### "):
            p = document.add_paragraph(style="Heading 3")
            add_inline(p, stripped[5:])
            i += 1
            continue
        if stripped.startswith("> "):
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            set_cell_shading(cell, PALE_GOLD)
            set_cell_margins(cell, 100, 140, 100, 140)
            p = cell.paragraphs[0]
            add_inline(p, stripped[2:], default_size=9.5)
            i += 1
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if bullet:
            p = document.add_paragraph(style="List Bullet")
            add_inline(p, bullet.group(1))
            i += 1
            continue
        if numbered:
            # Preserve the explicit Markdown number. Word's built-in List Number
            # style otherwise continues numbering across unrelated lists and can
            # turn a new 1–3 list into 11–13 or references 1–20 into 35–54.
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.65)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            prefix = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(prefix, size=10.5)
            add_inline(p, numbered.group(2))
            i += 1
            continue
        p = document.add_paragraph()
        p.paragraph_format.keep_together = False
        add_inline(p, stripped)
        i += 1


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    setup_styles(document)
    for section in document.sections:
        setup_section(section)
    document.core_properties.title = "LegumeGenomeFM：研究设计、数据合同与基准评测方案"
    document.core_properties.subject = "豆科超长上下文基因组基础模型的正式研究设计与证据边界"
    document.core_properties.author = "LegumeGenomeFM Project"
    document.core_properties.keywords = "legume genome foundation model, HierMamba, 256K, AgroNT, PlantCAD2, Evo 2"
    document.core_properties.comments = "Generated from LEGUMEGENOMEFM_RESEARCH_DESIGN.md; performance not yet run."
    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    add_cover(document)
    add_toc(document)
    render_markdown(document, text)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
